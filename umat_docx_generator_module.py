"""
UMAT DOCX Generator - Creates Word documents in exact UMAT format
Matches the format from sample files: acid___base.docx and work.docx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from typing import Dict, Any
import io

class UMATDocxGenerator:
    def __init__(self):
        pass
    
    def _format_date_umat_style(self, date_str: str) -> str:
        """Convert date to UMAT format: 12TH FEBRUARY, 2026"""
        try:
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            day = date_obj.day
            # Add ordinal suffix
            if 11 <= day <= 13:
                suffix = 'TH'
            else:
                suffix = {1: 'ST', 2: 'ND', 3: 'RD'}.get(day % 10, 'TH')
            return date_obj.strftime(f'%d{suffix} %B, %Y').upper()
        except:
            return date_str.upper()
    
    def generate_lab_report_docx(self, report_data: Dict[str, Any]) -> bytes:
        """
        Generate a complete lab report in UMAT format
        
        Args:
            report_data: Dictionary containing all report information
        
        Returns:
            DOCX file as bytes
        """
        
        doc = Document()
        
        # Set document margins (1 inch all around)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # ===== TITLE PAGE (All Centered) =====
        self._add_title_page(doc, report_data)
        
        # Page break after title page
        doc.add_page_break()
        
        # ===== CHAPTER 1: INTRODUCTION =====
        self._add_chapter_heading(doc, "CHAPTER 1")
        self._add_section_heading(doc, "INTRODUCTION")
        
        if 'introduction' in report_data.get('sections', {}):
            self._add_justified_paragraphs(doc, report_data['sections']['introduction'])
        
        # ===== CHAPTER 2: METHODOLOGY =====
        doc.add_page_break()
        self._add_chapter_heading(doc, "CHAPTER 2")
        self._add_section_heading(doc, "METHODOLOGY")
        
        # Equipment/Materials subsection
        if 'materials' in report_data.get('sections', {}):
            self._add_section_heading(doc, "Equipment/Materials", level=2)
            self._add_materials_list(doc, report_data['sections']['materials'])
        
        # Procedure
        if 'procedure' in report_data.get('sections', {}):
            self._add_justified_paragraphs(doc, report_data['sections']['procedure'])
        
        # ===== CHAPTER 3: DATA =====
        doc.add_page_break()
        self._add_chapter_heading(doc, "CHAPTER 3")
        self._add_section_heading(doc, "DATA")
        
        if 'data' in report_data.get('sections', {}):
            self._add_justified_paragraphs(doc, report_data['sections']['data'])
        else:
            p = doc.add_paragraph("[Data tables and calculations would be inserted here]")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # ===== CHAPTER 4: RESULTS AND DISCUSSION =====
        doc.add_page_break()
        self._add_chapter_heading(doc, "CHAPTER 4")
        self._add_section_heading(doc, "RESULTS AND DISCUSSION")
        
        if 'discussion' in report_data.get('sections', {}):
            self._add_justified_paragraphs(doc, report_data['sections']['discussion'])
        
        # Conclusion subsection
        if 'conclusion' in report_data.get('sections', {}):
            self._add_section_heading(doc, "Conclusion", level=2)
            self._add_justified_paragraphs(doc, report_data['sections']['conclusion'])
        
        # ===== REFERENCES =====
        doc.add_page_break()
        self._add_section_heading(doc, "REFERENCES")
        
        if 'references' in report_data.get('sections', {}):
            self._add_references(doc, report_data['sections']['references'])
        
        # Convert to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    def _add_title_page(self, doc: Document, report_data: Dict):
        """Add UMAT-formatted title page (all centered)"""
        
        # University details
        self._add_centered_paragraph(doc, "UNIVERSITY OF MINES AND TECHNOLOGY, (UMAT).")
        self._add_centered_paragraph(doc, "TARKWA, GHANA.")
        self._add_centered_paragraph(doc, "SCHOOL OF PETROLEUM STUDIES.")
        self._add_centered_paragraph(doc, "DEPARTMENT OF CHEMICAL AND PETROCHEMICAL ENGINEERING")
        
        # Space for logo
        doc.add_paragraph()
        self._add_centered_paragraph(doc, "[UMAT LOGO HERE]")
        doc.add_paragraph()
        
        # Report title
        title = report_data.get('title', '').upper()
        self._add_centered_paragraph(doc, f"A REPORT ON {title}.")
        doc.add_paragraph()
        
        # Student details
        student_name = report_data.get('student_name', '[STUDENT NAME]').upper()
        self._add_centered_paragraph(doc, f"BY: {student_name}.")
        
        student_id = report_data.get('student_id', '[STUDENT ID]')
        self._add_centered_paragraph(doc, student_id)
        
        course_code = report_data.get('course_code', 'CH 273')
        self._add_centered_paragraph(doc, f"{course_code}.")
        
        group_number = report_data.get('group_number', '')
        if group_number:
            self._add_centered_paragraph(doc, f"GROUP {group_number}")
        
        course_name = report_data.get('course_name', 'CHEMISTRY LABORATORY PRACTICE').upper()
        self._add_centered_paragraph(doc, course_name)
        
        lecturer = report_data.get('lecturer', '[COURSE LECTURER]').upper()
        self._add_centered_paragraph(doc, f"COURSE LECTURER: {lecturer}.")
        
        # Date
        date_str = report_data.get('date', datetime.now().strftime('%Y-%m-%d'))
        formatted_date = self._format_date_umat_style(date_str)
        self._add_centered_paragraph(doc, formatted_date)
    
    def _add_centered_paragraph(self, doc: Document, text: str):
        """Add a centered paragraph"""
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_chapter_heading(self, doc: Document, text: str):
        """Add CHAPTER heading (Heading 1)"""
        p = doc.add_heading(text, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_section_heading(self, doc: Document, text: str, level: int = 2):
        """Add section heading (Heading 2)"""
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_justified_paragraphs(self, doc: Document, content: str):
        """Add justified paragraphs from content"""
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _add_materials_list(self, doc: Document, materials: str):
        """Add materials as a list"""
        items = materials.split('\n') if isinstance(materials, str) else materials
        for item in items:
            if item.strip():
                p = doc.add_paragraph(item.strip(), style='List Bullet')
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _add_references(self, doc: Document, references: str):
        """Add references section"""
        refs = references.split('\n\n') if '\n\n' in references else references.split('\n')
        for ref in refs:
            if ref.strip() and not ref.strip().startswith('*Note'):
                p = doc.add_paragraph(ref.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
